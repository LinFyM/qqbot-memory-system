# -*- coding: utf-8 -*-
from typing import Optional, Callable, List, Tuple
import os
import subprocess

Logger = object

def extract_text_and_images_from_file(local_path: str,
                                      output_image_dir: str,
                                      metrics_add: Optional[Callable[[str, int], None]] = None,
                                      logger: Optional[Logger] = None) -> Tuple[str, List[str]]:
    """
    从文件中提取文本和图片
    
    Args:
        local_path: 本地文件路径
        output_image_dir: 提取的图片保存目录
        metrics_add: 指标记录函数
        logger: 日志记录器
    
    Returns:
        (文本内容, 提取的图片路径列表)
    """
    text = ""
    image_paths: List[str] = []
    
    try:
        import mimetypes
        _, ext = os.path.splitext(local_path.lower())
        mtype = mimetypes.guess_type(local_path)[0] or ""
        if logger: logger.info(f"尝试抽取文件正文和图片: {local_path} ({mtype or ext})")
        
        # PDF文件处理
        if ext in {".pdf"} or "pdf" in mtype:
            # 提取文本
            try:
                import pdfminer.high_level as pdfminer
                text = pdfminer.extract_text(local_path) or ""
            except Exception:
                if logger: logger.warning("pdfminer 不可用，尝试 pdftotext")
                try:
                    result = subprocess.run(["pdftotext", local_path, "-"], capture_output=True, check=True)
                    text = result.stdout.decode("utf-8", errors="ignore")
                except Exception:
                    if logger: logger.warning("pdftotext 不可用，放弃PDF文本抽取")
            
            # 提取图片
            try:
                from pdf2image import convert_from_path
                from datetime import datetime
                from uuid import uuid4
                images = convert_from_path(local_path, dpi=200)
                os.makedirs(output_image_dir, exist_ok=True)
                for i, img in enumerate(images):
                    filename = f"{datetime.utcnow().strftime('%Y%m%d%H%M%S%f')}_{uuid4().hex}_{i}.png"
                    img_path = os.path.join(output_image_dir, filename)
                    img.save(img_path, "PNG")
                    image_paths.append(img_path)
                if logger and image_paths: logger.info(f"✅ 从PDF中提取到 {len(image_paths)} 张图片")
            except Exception as img_err:
                if logger: logger.warning(f"PDF图片提取失败: {img_err}")
        
        # DOCX文件处理
        elif ext in {".docx"} or "officedocument.wordprocessingml.document" in mtype:
            try:
                import docx
                from datetime import datetime
                from uuid import uuid4
                doc = docx.Document(local_path)
                # 提取文本
                text = "\n".join(p.text for p in doc.paragraphs)
                
                # 提取图片
                os.makedirs(output_image_dir, exist_ok=True)
                for i, rel in enumerate(doc.part.rels.values()):
                    if "image" in rel.target_ref:
                        try:
                            image_part = rel.target_part
                            image_data = image_part.blob
                            # 根据图片类型确定扩展名
                            ext_map = {
                                "image/png": ".png",
                                "image/jpeg": ".jpg",
                                "image/jpg": ".jpg",
                                "image/gif": ".gif",
                                "image/bmp": ".bmp"
                            }
                            content_type = image_part.content_type
                            img_ext = ext_map.get(content_type, ".png")
                            filename = f"{datetime.utcnow().strftime('%Y%m%d%H%M%S%f')}_{uuid4().hex}_{i}{img_ext}"
                            img_path = os.path.join(output_image_dir, filename)
                            with open(img_path, "wb") as f:
                                f.write(image_data)
                            image_paths.append(img_path)
                        except Exception as img_err:
                            if logger: logger.warning(f"DOCX图片提取失败 (第{i}张): {img_err}")
                if logger and image_paths: logger.info(f"✅ 从DOCX中提取到 {len(image_paths)} 张图片")
            except Exception:
                if logger: logger.warning("python-docx 不可用，跳过DOCX抽取")
        
        # 简单文本文件
        elif ext in {".txt", ".md", ".log"} or mtype.startswith("text/"):
            with open(local_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        # 其他文件类型尝试OCR
        else:
            try:
                from PIL import Image
                import pytesseract
                text = pytesseract.image_to_string(Image.open(local_path))
            except Exception:
                if logger: logger.info("OCR 不可用或非图片，跳过")
        
        text = text.strip() if text else ""
        if text and metrics_add: metrics_add("file_extract_ok", 1)
        if not text and metrics_add: metrics_add("file_extract_empty", 1)
        if image_paths and metrics_add: metrics_add("file_image_extract_ok", len(image_paths))
        return text, image_paths
        
    except Exception as e:
        if logger: logger.warning(f"⚠️ 文件抽取失败: {e}")
        if metrics_add: metrics_add("file_extract_fail", 1)
        return "", []

def extract_text_from_file(local_path: str,
                           metrics_add: Optional[Callable[[str, int], None]] = None,
                           logger: Optional[Logger] = None) -> str:
    text_chunks = []
    try:
        import mimetypes
        _, ext = os.path.splitext(local_path.lower())
        mtype = mimetypes.guess_type(local_path)[0] or ""
        if logger: logger.info(f"尝试抽取文件正文: {local_path} ({mtype or ext})")
        # 简单类型
        if ext in {".txt", ".md", ".log"} or mtype.startswith("text/"):
            with open(local_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            text_chunks.append(text)
        elif ext in {".pdf"} or "pdf" in mtype:
            try:
                import pdfminer.high_level as pdfminer
                text = pdfminer.extract_text(local_path) or ""
                text_chunks.append(text)
            except Exception:
                if logger: logger.warning("pdfminer 不可用，尝试 pdftotext")
                try:
                    result = subprocess.run(["pdftotext", local_path, "-"], capture_output=True, check=True)
                    text_chunks.append(result.stdout.decode("utf-8", errors="ignore"))
                except Exception:
                    if logger: logger.warning("pdftotext 不可用，放弃PDF抽取")
        elif ext in {".docx"} or "officedocument.wordprocessingml.document" in mtype:
            try:
                import docx
                doc = docx.Document(local_path)
                text = "\n".join(p.text for p in doc.paragraphs)
                text_chunks.append(text)
            except Exception:
                if logger: logger.warning("python-docx 不可用，跳过DOCX抽取")
        else:
            # 兜底：尝试OCR（可选）
            try:
                from PIL import Image
                import pytesseract
                text = pytesseract.image_to_string(Image.open(local_path))
                text_chunks.append(text)
            except Exception:
                if logger: logger.info("OCR 不可用或非图片，跳过")
        text = "\n".join([t for t in text_chunks if t]).strip()
        if text and metrics_add: metrics_add("file_extract_ok", 1)
        if not text and metrics_add: metrics_add("file_extract_empty", 1)
        return text
    except Exception as e:
        if logger: logger.warning(f"⚠️ 文件抽取失败: {e}")
        if metrics_add: metrics_add("file_extract_fail", 1)
        return ""

def download_and_extract_webpage(url: str,
                                 metrics_add: Optional[Callable[[str, int], None]] = None,
                                 logger: Optional[Logger] = None,
                                 headers: Optional[dict] = None) -> str:
    try:
        import requests
        if logger: logger.info(f"抓取网页: {url}")
        req_headers = headers or {}
        resp = requests.get(url, timeout=10, headers=req_headers)
        resp.raise_for_status()
        html = resp.text or ""
        text = ""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            for s in soup(["script", "style", "noscript"]):
                s.decompose()
            text = (soup.get_text("\n") or "").strip()
        except Exception:
            if logger: logger.warning("BeautifulSoup 不可用，使用正则粗提取")
            import re
            text = re.sub(r"<[^>]+>", " ", html)
            text = re.sub(r"\s+", " ", text).strip()
        if metrics_add: metrics_add("web_extract_ok", 1)
        return text
    except Exception as e:
        if logger: logger.warning(f"⚠️ 网页抓取失败: {e}")
        if metrics_add: metrics_add("web_extract_fail", 1)
        return ""


